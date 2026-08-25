"""
app/ai/document_processor.py
Admin Document Processing and RAG Ingestion Engine.
Handles:
- Text extraction from PDF (using pypdf) and DOCX (using python-docx)
- Text cleaning and normalization
- Semantic chunking with page number and audit metadata preservation
- Vector embedding generation using the existing embedding model
- Direct indexing and re-indexing into the existing ChromaDB / In-Memory vector store
- Removal/deletion of chunks and embeddings on document delete/replace
- Error handling (empty PDF, corrupt file, invalid format, scanned document detection)
"""
import os
import re
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional
from langchain_core.documents import Document

from app.ai.chunker import chunk_documents
from app.ai.embeddings import get_embedding_model
from app.ai.vector_store import (
    build_or_load_vector_store,
    add_documents_to_vector_store,
    delete_documents_from_vector_store,
)


def calculate_file_hash(file_path: str) -> str:
    """Calculates SHA256 hash of a file for duplicate detection and fingerprinting."""
    if not os.path.exists(file_path):
        return ""
    hasher = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def clean_extracted_text(text: str) -> str:
    """Normalizes whitespace and removes unprintable or null characters."""
    if not text:
        return ""
    # Replace null bytes
    text = text.replace("\x00", "")
    # Normalize excessive line breaks and whitespace
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


def extract_text_from_pdf(file_path: str) -> Tuple[bool, str, List[Dict[str, Any]], Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF document using pypdf.
    Returns: (success, error_message, pages_data, metadata)
    """
    if not os.path.exists(file_path):
        return False, "File does not exist on server.", [], {}

    pages_data = []
    metadata = {
        "page_count": 0,
        "total_characters": 0,
        "is_scanned": False
    }

    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)

        if reader.is_encrypted:
            try:
                # Try decrypting with empty password
                reader.decrypt("")
            except Exception:
                return False, "PDF is encrypted and password protected.", [], {}

        page_count = len(reader.pages)
        metadata["page_count"] = page_count

        if page_count == 0:
            return False, "PDF document contains 0 pages.", [], metadata

        total_chars = 0
        for page_idx, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                page_text = clean_extracted_text(page_text)
            except Exception as pe:
                page_text = ""

            if page_text:
                total_chars += len(page_text)
                pages_data.append({
                    "page_number": page_idx + 1,
                    "text": page_text
                })

        metadata["total_characters"] = total_chars

        # Check if text was extracted or if it is an image-only / empty scanned PDF
        if total_chars < 20:
            metadata["is_scanned"] = True
            return (
                False, 
                "Unable to extract text from PDF (document may be empty or contain only scanned images).", 
                [], 
                metadata
            )

        return True, "Text extracted successfully.", pages_data, metadata

    except Exception as e:
        return False, f"Failed to parse PDF file: {str(e)}", [], metadata


def extract_text_from_docx(file_path: str) -> Tuple[bool, str, List[Dict[str, Any]], Dict[str, Any]]:
    """
    Extracts text paragraphs and tables from a Word DOCX document using python-docx.
    Returns: (success, error_message, pages_data, metadata)
    """
    if not os.path.exists(file_path):
        return False, "File does not exist on server.", [], {}

    pages_data = []
    metadata = {
        "page_count": 1,
        "total_characters": 0,
        "is_scanned": False
    }

    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs_text = []

        for p in doc.paragraphs:
            text = clean_extracted_text(p.text)
            if text:
                paragraphs_text.append(text)

        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_cells = [clean_extracted_text(cell.text) for cell in row.cells if clean_extracted_text(cell.text)]
                if row_cells:
                    paragraphs_text.append(" | ".join(row_cells))

        full_text = "\n\n".join(paragraphs_text)
        metadata["total_characters"] = len(full_text)

        if not full_text or len(full_text) < 10:
            return False, "DOCX document is empty or contains no readable text.", [], metadata

        pages_data.append({
            "page_number": 1,
            "text": full_text
        })

        return True, "Text extracted successfully.", pages_data, metadata

    except Exception as e:
        return False, f"Failed to parse DOCX file: {str(e)}", [], metadata


def extract_document_text(file_path: str, file_type: str = "pdf") -> Tuple[bool, str, List[Dict[str, Any]], Dict[str, Any]]:
    """Universal dispatcher for PDF and DOCX text extraction."""
    ext = os.path.splitext(file_path)[1].lower().replace('.', '')
    
    if ext == 'pdf' or file_type == 'application/pdf' or file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ('docx', 'doc') or 'word' in str(file_type).lower():
        return extract_text_from_docx(file_path)
    else:
        return False, f"Unsupported file extension '.{ext}'. Supported types: PDF, DOCX.", [], {}


def build_chunks_from_pages(
    pages_data: List[Dict[str, Any]], 
    doc_metadata: Dict[str, Any],
    chunk_size: int = 450, 
    chunk_overlap: int = 50
) -> List[Document]:
    """
    Transforms extracted page text into LangChain Document chunks,
    attaching complete source metadata to every generated chunk.
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""]
    )

    doc_id = str(doc_metadata.get("document_id") or doc_metadata.get("id", ""))
    doc_name = str(doc_metadata.get("name") or doc_metadata.get("title") or doc_metadata.get("file_name", "Document"))
    source_filename = str(doc_metadata.get("file_name") or doc_metadata.get("original_filename") or doc_name)
    category = str(doc_metadata.get("category", "Academic"))
    department = str(doc_metadata.get("department", "All"))
    source_type = str(doc_metadata.get("source_type", "admin_document"))
    version = int(doc_metadata.get("version", 1))
    uploaded_by = str(doc_metadata.get("uploaded_by", "admin"))
    uploaded_at = str(doc_metadata.get("uploaded_at") or doc_metadata.get("created_at") or datetime.utcnow().isoformat())

    all_chunks = []
    global_chunk_idx = 0

    for page in pages_data:
        page_num = page.get("page_number", 1)
        raw_text = page.get("text", "")
        if not raw_text.strip():
            continue

        # Split page text into chunks
        page_split_texts = text_splitter.split_text(raw_text)

        for chunk_text in page_split_texts:
            clean_chunk = clean_extracted_text(chunk_text)
            if not clean_chunk:
                continue

            chunk_meta = {
                "source": source_filename,
                "document_id": doc_id,
                "document_name": doc_name,
                "category": category,
                "department": department,
                "source_type": source_type,
                "page_number": page_num,
                "chunk_index": global_chunk_idx,
                "version": version,
                "uploaded_by": uploaded_by,
                "uploaded_at": uploaded_at,
                "is_active": True,
                "file_url": doc_metadata.get("file_url", ""),
            }

            all_chunks.append(Document(page_content=clean_chunk, metadata=chunk_meta))
            global_chunk_idx += 1

    return all_chunks


def process_and_index_document(
    document_id: str,
    file_path: str,
    doc_metadata: Dict[str, Any],
    vector_store = None
) -> Tuple[bool, str, int, Dict[str, Any]]:
    """
    Full pipeline execution for a single document:
    1. Extract Text
    2. Clean Text
    3. Split into Semantic Chunks
    4. Generate Embeddings & Store in existing Vector Store
    5. Invalidate vector search caches

    Returns: (success, error_message, chunk_count, stats_dict)
    """
    if not vector_store:
        from app.ai.rag_pipeline import get_rag_pipeline
        pipeline = get_rag_pipeline()
        vector_store = pipeline.vector_store

    doc_metadata["document_id"] = document_id
    file_type = doc_metadata.get("file_type", "pdf")

    # Step 1: Extract Text
    success, err_msg, pages_data, extract_meta = extract_document_text(file_path, file_type=file_type)
    if not success:
        return False, err_msg, 0, extract_meta

    # Step 2 & 3: Chunking with source metadata
    chunks = build_chunks_from_pages(pages_data, doc_metadata)
    if not chunks:
        return False, "No usable text chunks could be generated from document.", 0, extract_meta

    # Step 4: First remove any previous chunks for this document (prevent duplicates)
    delete_documents_from_vector_store(vector_store, document_id)

    # Step 5: Add chunks to existing Vector Store
    try:
        add_documents_to_vector_store(vector_store, chunks)
    except Exception as e:
        return False, f"Vector store embedding/indexing error: {str(e)}", 0, extract_meta

    # Step 6: Invalidate query/retrieval memory caches
    _clear_rag_caches()

    stats = {
        "page_count": extract_meta.get("page_count", 1),
        "chunk_count": len(chunks),
        "total_characters": extract_meta.get("total_characters", 0),
        "file_hash": calculate_file_hash(file_path),
        "indexed_at": datetime.utcnow().isoformat()
    }

    return True, f"Successfully indexed {len(chunks)} chunks across {stats['page_count']} pages.", len(chunks), stats


def remove_document_from_rag(document_id: str, vector_store = None) -> bool:
    """Removes all vector embeddings and chunks associated with a document_id."""
    if not vector_store:
        from app.ai.rag_pipeline import get_rag_pipeline
        pipeline = get_rag_pipeline()
        vector_store = pipeline.vector_store

    deleted = delete_documents_from_vector_store(vector_store, document_id)
    _clear_rag_caches()
    return deleted


def _clear_rag_caches():
    """Flushes LRU caches in retriever and rag_pipeline so updated index is queried immediately."""
    try:
        from app.ai.retriever import _VECTOR_CACHE
        _VECTOR_CACHE.clear()
    except Exception:
        pass

    try:
        from app.ai.rag_pipeline import _RESPONSE_CACHE
        _RESPONSE_CACHE.clear()
    except Exception:
        pass
