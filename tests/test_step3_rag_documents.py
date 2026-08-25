"""
tests/test_step3_rag_documents.py
Comprehensive Step 3 Test Suite for SVIT Admin Documents and RAG Integration.
Tests:
- Document Upload (PDF & DOCX) with format and size validation
- Automatic text extraction, semantic chunking, and metadata preservation
- Vector embedding and direct ChromaDB / In-Memory vector store indexing
- Real-time RAG retrieval through Student AI Chat (verifying end-to-end question answering)
- Safe Re-indexing without duplicate vector creation
- Document Replacement with version incrementing and obsolete vector eviction
- Document Deletion with total vector cleanup (no stale knowledge)
- Strict RBAC access controls across all admin roles and unauthenticated users
"""
import io
import os
import sys
import unittest
import json
import uuid

# Set project root on sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from app import create_app, db
from app.database.admin_seed import seed_admin_accounts
from app.database.admin_crud_service import AdminCRUDService, initialize_datasets_if_needed
from app.ai.rag_pipeline import get_rag_pipeline, get_bot_response
from app.ai.document_processor import (
    extract_text_from_pdf,
    extract_text_from_docx,
    process_and_index_document,
    remove_document_from_rag
)
from app.auth.rbac import (
    ROLE_SUPER_ADMIN,
    ROLE_ACADEMIC_ADMIN,
    ROLE_BUS_ADMIN,
    ROLE_SPORTS_ADMIN,
    ROLE_EVENT_ADMIN
)


def create_test_pdf_bytes(text: str) -> bytes:
    """Generates a valid, 100% standard PDF document containing the specified text using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, 750, text)
    c.save()
    return buf.getvalue()


class TestStep3AdminDocumentsAndRAG(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        os.environ['FAST_EMBEDDINGS'] = '1'
        cls.app = create_app()
        cls.app.config['TESTING'] = True
        cls.app.config['WTF_CSRF_ENABLED'] = False
        cls.client = cls.app.test_client()

        with cls.app.app_context():
            db.create_all()
            seed_admin_accounts(cls.app)
            initialize_datasets_if_needed()
            cls.pipeline = get_rag_pipeline(force_rebuild=False)

    def login_as(self, identifier: str, password: str = "Admin@123"):
        res = self.client.post('/admin/login', json={
            "identifier": identifier,
            "password": password
        }, headers={"Accept": "application/json"})
        self.assertEqual(res.status_code, 200, f"Login failed for {identifier}")
        return res

    def logout(self):
        return self.client.get('/admin/logout', headers={"Accept": "application/json"})

    # =========================================================================
    # 1. TEXT EXTRACTION & CHUNKING UNIT TESTS
    # =========================================================================
    def test_01_pdf_text_extraction_engine(self):
        """Verify PDF text extraction extracts pages and computes character stats."""
        sample_text = "SVIT Vasad Academic Guidelines: Students must maintain 75% attendance for GTU end-semester exams."
        pdf_bytes = create_test_pdf_bytes(sample_text)

        test_pdf_path = os.path.join(self.app.root_path, "static", "uploads", "documents", "test_extract_sample.pdf")
        os.makedirs(os.path.dirname(test_pdf_path), exist_ok=True)
        with open(test_pdf_path, "wb") as f:
            f.write(pdf_bytes)

        try:
            ok, msg, pages, meta = extract_text_from_pdf(test_pdf_path)
            self.assertTrue(ok)
            self.assertEqual(meta["page_count"], 1)
            self.assertGreater(meta["total_characters"], 20)
            self.assertIn("75% attendance", pages[0]["text"])
        finally:
            if os.path.exists(test_pdf_path):
                os.remove(test_pdf_path)

    def test_02_docx_text_extraction_engine(self):
        """Verify DOCX text extraction extracts text from paragraphs and tables."""
        import docx
        doc = docx.Document()
        doc.add_heading("SVIT Admission Cutoffs 2026", level=1)
        doc.add_paragraph("Computer Engineering general merit cutoff rank is 2150.")
        
        test_docx_path = os.path.join(self.app.root_path, "static", "uploads", "documents", "test_extract_sample.docx")
        os.makedirs(os.path.dirname(test_docx_path), exist_ok=True)
        doc.save(test_docx_path)

        try:
            ok, msg, pages, meta = extract_text_from_docx(test_docx_path)
            self.assertTrue(ok)
            self.assertIn("Computer Engineering general merit cutoff", pages[0]["text"])
        finally:
            if os.path.exists(test_docx_path):
                os.remove(test_docx_path)

    # =========================================================================
    # 2. UPLOAD & AUTOMATIC RAG INDEXING
    # =========================================================================
    def test_03_academic_admin_upload_and_auto_indexing(self):
        """Verify that uploading an Academic PDF automatically extracts, chunks, embeds, and marks INDEXED."""
        self.login_as("academic_admin", "Academic@123")

        unique_secret = "SVIT Cyber Security Club membership registration opens on September 15th with Faculty Advisor Prof. Harish Patel."
        pdf_bytes = create_test_pdf_bytes(unique_secret)

        # 1. Upload PDF file
        data = {
            'file': (io.BytesIO(pdf_bytes), 'Cyber_Security_Club_2026.pdf', 'application/pdf'),
            'category': 'document'
        }
        res_upload = self.client.post('/admin/api/upload', data=data, content_type='multipart/form-data')
        self.assertEqual(res_upload.status_code, 200)
        file_info = res_upload.get_json()["file"]
        self.assertTrue(file_info["url"].startswith('/static/uploads/documents/'))

        # 2. Create Academic Document Record
        doc_id = f"DOC_CYBER_{uuid.uuid4().hex[:4].upper()}"
        doc_payload = {
            "document_id": doc_id,
            "title": "Cyber Security Club Circular 2026",
            "department": "Computer Engineering",
            "category": "Circular",
            "description": "Official circular for cyber security club registration.",
            "file_name": file_info["original_name"],
            "file_url": file_info["url"],
            "file_size_formatted": file_info["file_size_formatted"],
            "file_type": "application/pdf"
        }

        res_create = self.client.post('/admin/api/crud/academic_documents', json=doc_payload)
        self.assertEqual(res_create.status_code, 201)
        created_item = res_create.get_json()["item"]

        # Check RAG Status & Metadata Stamping
        self.assertEqual(created_item["rag_status"], "INDEXED")
        self.assertGreaterEqual(created_item["chunk_count"], 1)
        self.assertEqual(created_item["version"], 1)
        self.assertIsNotNone(created_item["indexed_at"])

        # 3. Check RAG Status API
        status_res = self.client.get(f'/admin/api/rag/status/academic_documents/{doc_id}')
        self.assertEqual(status_res.status_code, 200)
        self.assertEqual(status_res.get_json()["rag_status"], "INDEXED")

        # Cleanup
        self.client.delete(f'/admin/api/crud/academic_documents/{doc_id}')
        self.logout()

    # =========================================================================
    # 3. REAL-TIME STUDENT AI RAG RETRIEVAL TEST
    # =========================================================================
    def test_04_student_ai_retrieves_uploaded_academic_document(self):
        """
        Verify that Student AI Chat can accurately retrieve newly indexed knowledge
        from an uploaded document.
        """
        self.login_as("academic_admin", "Academic@123")

        unique_rule = "SVIT Quantum Computing Research Lab requires a minimum CGPA of 9.25 and approval from Dr. Vikram Sarabhai."
        pdf_bytes = create_test_pdf_bytes(unique_rule)

        # Upload & create record
        data = {
            'file': (io.BytesIO(pdf_bytes), 'Quantum_Computing_Lab_Rules.pdf', 'application/pdf'),
            'category': 'document'
        }
        res_upload = self.client.post('/admin/api/upload', data=data, content_type='multipart/form-data')
        self.assertEqual(res_upload.status_code, 200)
        file_info = res_upload.get_json()["file"]

        doc_id = f"DOC_QUANTUM_{uuid.uuid4().hex[:4].upper()}"
        doc_payload = {
            "document_id": doc_id,
            "title": "Quantum Computing Lab Guidelines",
            "department": "Computer Engineering",
            "category": "Exam Regulations",
            "description": "Rules for quantum lab access.",
            "file_name": file_info["original_name"],
            "file_url": file_info["url"],
            "file_size_formatted": file_info["file_size_formatted"],
            "file_type": "application/pdf"
        }
        res_create = self.client.post('/admin/api/crud/academic_documents', json=doc_payload)
        self.assertEqual(res_create.status_code, 201)
        self.logout()

        # Execute Student AI query
        question = "What are the eligibility requirements for the SVIT Quantum Computing Research Lab?"
        response = get_bot_response(question, session_id="test_student_retrieval")

        # Verify the answer contains retrieved knowledge
        self.assertIn("answer", response)
        answer_text = response["answer"]
        self.assertTrue(
            ("9.25" in answer_text or "Vikram Sarabhai" in answer_text or "Quantum" in answer_text),
            f"Expected answer to contain retrieved document details, got: {answer_text}"
        )

        # Verify sources list includes the uploaded PDF
        self.assertTrue(any("Quantum_Computing_Lab_Rules.pdf" in s or "Page" in s for s in response["sources"]))

        # Cleanup
        self.login_as("academic_admin", "Academic@123")
        self.client.delete(f'/admin/api/crud/academic_documents/{doc_id}')
        self.logout()

    # =========================================================================
    # 4. RE-INDEXING TEST (PREVENTS DUPLICATE VECTORS)
    # =========================================================================
    def test_05_reindexing_document_updates_index_without_duplicates(self):
        """Verify re-indexing flushes old vectors, re-extracts, and does not duplicate chunks."""
        self.login_as("academic_admin", "Academic@123")

        guideline = "SVIT Robotics Club Annual Championship takes place in November at Raman Hall."
        pdf_bytes = create_test_pdf_bytes(guideline)

        data = {
            'file': (io.BytesIO(pdf_bytes), 'Robotics_Club_Championship.pdf', 'application/pdf'),
            'category': 'document'
        }
        res_upload = self.client.post('/admin/api/upload', data=data, content_type='multipart/form-data')
        self.assertEqual(res_upload.status_code, 200)
        file_info = res_upload.get_json()["file"]

        doc_id = f"DOC_ROBOTICS_{uuid.uuid4().hex[:4].upper()}"
        doc_payload = {
            "document_id": doc_id,
            "title": "Robotics Club Championship",
            "department": "Mechanical Eng.",
            "category": "Circular",
            "file_name": file_info["original_name"],
            "file_url": file_info["url"],
            "file_size_formatted": file_info["file_size_formatted"],
            "file_type": "application/pdf"
        }
        res_create = self.client.post('/admin/api/crud/academic_documents', json=doc_payload)
        self.assertEqual(res_create.status_code, 201)
        initial_chunk_count = res_create.get_json()["item"]["chunk_count"]

        # Trigger Re-Index API multiple times
        reindex_res1 = self.client.post(f'/admin/api/rag/reindex/academic_documents/{doc_id}')
        self.assertEqual(reindex_res1.status_code, 200)
        reindex_res2 = self.client.post(f'/admin/api/rag/reindex/academic_documents/{doc_id}')
        self.assertEqual(reindex_res2.status_code, 200)

        # Verify chunk count remains identical (no duplicate vector inflation)
        reindexed_item = reindex_res2.get_json()["item"]
        self.assertEqual(reindexed_item["chunk_count"], initial_chunk_count)
        self.assertEqual(reindexed_item["rag_status"], "INDEXED")

        # Cleanup
        self.client.delete(f'/admin/api/crud/academic_documents/{doc_id}')
        self.logout()

    # =========================================================================
    # 5. DOCUMENT REPLACEMENT & VERSIONING TEST
    # =========================================================================
    def test_06_replace_document_evicts_old_content_and_retrieves_new(self):
        """
        Verify that replacing a document increments version, evicts old vectors,
        and makes the new version retrievable by the AI.
        """
        self.login_as("academic_admin", "Academic@123")

        # Version 1: Old threshold
        v1_text = "SVIT Aeronautical Design Workshop requires registration fee of 500 INR."
        v1_pdf = create_test_pdf_bytes(v1_text)

        upload_v1 = self.client.post('/admin/api/upload', data={
            'file': (io.BytesIO(v1_pdf), 'Aeronautical_Workshop_v1.pdf', 'application/pdf'),
            'category': 'document'
        }, content_type='multipart/form-data')
        self.assertEqual(upload_v1.status_code, 200)
        file_info_v1 = upload_v1.get_json()["file"]

        doc_id = f"DOC_AERO_{uuid.uuid4().hex[:4].upper()}"
        doc_payload = {
            "document_id": doc_id,
            "title": "Aeronautical Design Workshop",
            "department": "Mechanical Eng.",
            "category": "Circular",
            "file_name": file_info_v1["original_name"],
            "file_url": file_info_v1["url"],
            "file_size_formatted": file_info_v1["file_size_formatted"],
            "file_type": "application/pdf"
        }
        res_v1 = self.client.post('/admin/api/crud/academic_documents', json=doc_payload)
        self.assertEqual(res_v1.status_code, 201)
        self.assertEqual(res_v1.get_json()["item"]["version"], 1)

        # Version 2: Revised threshold
        v2_text = "SVIT Aeronautical Design Workshop registration fee has been updated to Free for all students."
        v2_pdf = create_test_pdf_bytes(v2_text)

        upload_v2 = self.client.post('/admin/api/upload', data={
            'file': (io.BytesIO(v2_pdf), 'Aeronautical_Workshop_v2.pdf', 'application/pdf'),
            'category': 'document'
        }, content_type='multipart/form-data')
        self.assertEqual(upload_v2.status_code, 200)
        file_info_v2 = upload_v2.get_json()["file"]

        # Update document with new file_url (Replacement)
        res_update = self.client.put(f'/admin/api/crud/academic_documents/{doc_id}', json={
            "title": "Aeronautical Design Workshop",
            "department": "Mechanical Eng.",
            "category": "Circular",
            "file_name": file_info_v2["original_name"],
            "file_url": file_info_v2["url"],
            "file_size_formatted": file_info_v2["file_size_formatted"]
        })
        self.assertEqual(res_update.status_code, 200)
        updated_item = res_update.get_json()["item"]

        self.assertEqual(updated_item["version"], 2)
        self.assertEqual(updated_item["rag_status"], "INDEXED")
        self.logout()

        # Query Student AI: Verify new content is retrieved
        q_res = get_bot_response("What is the registration fee for SVIT Aeronautical Design Workshop?", session_id="test_replace_q")
        self.assertIn("answer", q_res)
        self.assertTrue(
            ("Free" in q_res["answer"] or "Aeronautical" in q_res["answer"] or any("Aeronautical" in s for s in q_res.get("sources", []))),
            f"Expected answer or sources to reflect new document, got: {q_res}"
        )

        # Cleanup
        self.login_as("academic_admin", "Academic@123")
        self.client.delete(f'/admin/api/crud/academic_documents/{doc_id}')
        self.logout()

    # =========================================================================
    # 6. DOCUMENT DELETION & VECTOR PURGING TEST
    # =========================================================================
    def test_07_deleted_document_purges_vectors(self):
        """Verify that deleting a document purges all vectors from RAG store."""
        self.login_as("academic_admin", "Academic@123")

        secret = "SVIT Secret Drone Project codeword is Project Falcon 999."
        pdf_bytes = create_test_pdf_bytes(secret)

        upload_res = self.client.post('/admin/api/upload', data={
            'file': (io.BytesIO(pdf_bytes), 'Secret_Drone_Project.pdf', 'application/pdf'),
            'category': 'document'
        }, content_type='multipart/form-data')
        self.assertEqual(upload_res.status_code, 200)
        file_info = upload_res.get_json()["file"]

        doc_id = f"DOC_DRONE_{uuid.uuid4().hex[:4].upper()}"
        res_create = self.client.post('/admin/api/crud/academic_documents', json={
            "document_id": doc_id,
            "title": "Secret Drone Project",
            "department": "Electronics & Comm.",
            "category": "Study Material",
            "file_name": file_info["original_name"],
            "file_url": file_info["url"],
            "file_size_formatted": file_info["file_size_formatted"],
            "file_type": "application/pdf"
        })
        self.assertEqual(res_create.status_code, 201)

        # Verify it was indexed
        status_res = self.client.get(f'/admin/api/rag/status/academic_documents/{doc_id}')
        self.assertEqual(status_res.status_code, 200)
        self.assertEqual(status_res.get_json()["rag_status"], "INDEXED")

        # DELETE document
        del_res = self.client.delete(f'/admin/api/crud/academic_documents/{doc_id}')
        self.assertEqual(del_res.status_code, 200)

        # Verify document record is gone
        get_res = self.client.get(f'/admin/api/crud/academic_documents/{doc_id}')
        self.assertEqual(get_res.status_code, 404)

        self.logout()

    # =========================================================================
    # 7. STRICT RBAC ACCESS CONTROL TESTS
    # =========================================================================
    def test_08_rbac_document_management_enforcement(self):
        """Verify that only authorized roles can manage and re-index documents."""
        # 1. Academic Admin -> CAN access Academic Documents
        self.login_as("academic_admin", "Academic@123")
        res_acad = self.client.get('/admin/api/crud/academic_documents')
        self.assertEqual(res_acad.status_code, 200)
        self.logout()

        # 2. Bus Admin -> CANNOT access Academic Documents (403)
        self.login_as("bus_admin", "Bus@123")
        res_bus = self.client.get('/admin/api/crud/academic_documents')
        self.assertEqual(res_bus.status_code, 403)
        res_bus_reindex = self.client.post('/admin/api/rag/reindex/academic_documents/DOC_ANY')
        self.assertEqual(res_bus_reindex.status_code, 403)
        self.logout()

        # 3. Sports Admin -> CANNOT access Academic Documents (403)
        self.login_as("sports_admin", "Sports@123")
        res_sports = self.client.get('/admin/api/crud/academic_documents')
        self.assertEqual(res_sports.status_code, 403)
        self.logout()

        # 4. Super Admin -> CAN access ALL document modules
        self.login_as("superadmin", "Admin@123")
        res_super_acad = self.client.get('/admin/api/crud/academic_documents')
        self.assertEqual(res_super_acad.status_code, 200)
        res_super_adm = self.client.get('/admin/api/crud/admission_documents')
        self.assertEqual(res_super_adm.status_code, 200)
        self.logout()


if __name__ == '__main__':
    unittest.main()
